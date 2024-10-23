# helper.py

import logging
import os
import re
import time
import json
import requests
import spacy
from docx import Document
from pdf2docx import Converter
from dotenv import load_dotenv

load_dotenv()
EMAIL = os.getenv('EMAIL', 'default_email@example.com')
if not EMAIL:
    raise ValueError("Biến môi trường EMAIL chưa được thiết lập. Vui lòng thêm vào tệp .env.")

# Tải mô hình spaCy
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    from spacy.cli import download
    download("en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

def convert_pdf_to_docx(pdf_file, docx_file):
    """
    Chuyển đổi tệp PDF thành DOCX.

    Args:
        pdf_file (str): Đường dẫn tới tệp PDF.
        docx_file (str): Đường dẫn tới tệp DOCX sẽ được tạo.

    Returns:
        str: Đường dẫn tới tệp DOCX đã được tạo.
    """
    logging.info(f"Chuyển đổi {pdf_file} thành {docx_file}")
    try:
        cv = Converter(pdf_file)
        cv.convert(docx_file, start=0, end=None)
        cv.close()
        logging.info(f"Đã chuyển đổi {pdf_file} thành {docx_file}.")
        if os.path.exists(docx_file):
            return docx_file
        else:
            logging.error(f"Tệp DOCX {docx_file} không được tạo.")
            raise FileNotFoundError(f"Tệp DOCX {docx_file} không được tạo.")
    except Exception as e:
        logging.error(f"Lỗi khi chuyển đổi PDF sang DOCX: {e}")
        raise

def clean_text(text):
    """
    Làm sạch văn bản bằng cách loại bỏ các ký tự xuống dòng, tab, khoảng trắng dư thừa, và dấu gạch ngang không hợp lệ.

    Args:
        text (str): Văn bản cần làm sạch.

    Returns:
        str: Văn bản đã được làm sạch.
    """
    text = remove_hyphenation(text)
    text = re.sub(r'[\n\r\t]+', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def remove_hyphenation(text):
    """
    Loại bỏ các dấu gạch ngang được sử dụng để ngắt dòng mà không ảnh hưởng đến các dấu gạch ngang hợp lệ trong từ.

    Args:
        text (str): Văn bản cần xử lý.

    Returns:
        str: Văn bản đã loại bỏ dấu gạch ngang ngắt dòng.
    """
    text = re.sub(r'-\s*\n\s*', '', text)
    text = re.sub(r'(\w+)-\s*(\w+)', r'\1\2', text)
    return text

def detect_citation_format(text):
    """
    Phát hiện định dạng trích dẫn trong văn bản.

    Args:
        text (str): Văn bản cần kiểm tra.

    Returns:
        str: Định dạng trích dẫn ('APA', 'IEEE', 'Chicago' hoặc 'Unknown format').
    """
    apa_pattern = re.compile(r'\b[A-Z][a-z]+, [A-Z]\. (?:\d{4}|\(\d{4}\))')
    ieee_pattern = re.compile(r'\[\d+\]')
    chicago_pattern = re.compile(r'\b[A-Z][a-z]+, [A-Z][a-z]+\. \d{4}\.')
    if apa_pattern.search(text):
        return "APA"
    elif ieee_pattern.search(text):
        return "IEEE"
    elif chicago_pattern.search(text):
        return "Chicago"
    else:
        return "Unknown format"

def extract_text_from_docx(docx_path):
    """
    Trích xuất văn bản từ tệp DOCX.

    Args:
        docx_path (str): Đường dẫn tới tệp DOCX.

    Returns:
        str: Văn bản trích xuất từ tệp DOCX.
    """
    try:
        doc = Document(docx_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        logging.info(f"Đã trích xuất văn bản từ DOCX: {docx_path}")
        return full_text
    except Exception as e:
        logging.error(f"Lỗi khi trích xuất văn bản từ DOCX '{docx_path}': {e}")
        return ""

# helper.py

def extract_sentences_from_docx(docx_file, output_file, exclude_sections=None):
    """
    Trích xuất các câu từ tệp DOCX, có thể loại bỏ các phần không mong muốn.

    Args:
        docx_file (str): Đường dẫn tới tệp DOCX.
        output_file (str): Đường dẫn tới tệp văn bản để lưu các câu.
        exclude_sections (list): Danh sách các tiêu đề phần cần loại bỏ (ví dụ: ['References']).

    Returns:
        list: Danh sách các câu trích xuất từ tệp DOCX.
    """
    logging.info(f"Đang đọc tệp DOCX: {docx_file}")

    # Đọc nội dung từ tệp DOCX
    doc = Document(docx_file)

    # Khởi tạo biến
    text = ''
    exclude = False

    # Duyệt qua các đoạn văn
    for para in doc.paragraphs:
        para_text = para.text.strip()

        # Kiểm tra các phần cần loại bỏ bằng regex
        if exclude_sections:
            for section in exclude_sections:
                # Sử dụng regex để kiểm tra chính xác tiêu đề phần, có thể có thêm số trang
                if re.match(rf'^{re.escape(section)}\b', para_text, re.IGNORECASE):
                    exclude = True
                    logging.info(f"Đã bắt đầu loại bỏ phần: {section}")
                    break

        # Nếu đang trong phần loại bỏ, bỏ qua các đoạn văn
        if exclude:
            # Nếu gặp một tiêu đề mới không nằm trong exclude_sections, dừng loại bỏ
            is_new_section = False
            if exclude_sections:
                for section in exclude_sections:
                    if re.match(rf'^{re.escape(section)}\b', para_text, re.IGNORECASE):
                        is_new_section = True
                        break
            if is_new_section:
                continue  # Tiếp tục loại bỏ
            else:
                # Không phải phần loại bỏ, tiếp tục bao gồm văn bản
                exclude = False

        if not exclude:
            text += para_text + '\n'

    # Làm sạch và xử lý văn bản bằng spaCy
    cleaned_text = clean_text(text)
    spacy_doc = nlp(cleaned_text)

    # Trích xuất các câu
    sentences = []
    with open(output_file, "w", encoding="utf-8") as f:
        for sent in spacy_doc.sents:
            cleaned_sentence = sent.text.strip()
            f.write(cleaned_sentence + "\n")
            sentences.append(cleaned_sentence)

    logging.info(f"Đã lưu các câu vào tệp: {output_file}")
    return sentences

def search_paper_doi_with_info(reference):
    """
    Tìm kiếm DOI của bài báo và lấy thông tin chi tiết từ Crossref API.

    Args:
        reference (str): Chuỗi tham chiếu.

    Returns:
        tuple: (doi, crossref_info)
            - doi (str): DOI của bài báo.
            - crossref_info (dict): Thông tin chi tiết của bài báo.
    """
    api_url = "https://api.crossref.org/works"
    params = {
        'query.title': reference,
        'rows': 1
    }

    try:
        response = requests.get(api_url, params=params, timeout=10)
        if response.status_code == 200:
            data = response.json()
            items = data.get('message', {}).get('items', [])
            if items:
                doi = items[0].get('DOI', None)
                title = items[0].get('title', ['No Title'])[0]
                authors_list = items[0].get('author', [])
                authors = ', '.join([f"{author.get('given', '')} {author.get('family', '')}" for author in authors_list])
                year = items[0].get('issued', {}).get('date-parts', [[None]])[0][0]

                crossref_info = {
                    'title': title,
                    'authors': authors,
                    'year': str(year) if year else ''
                }

                logging.info(f"Tìm thấy bài báo: {title} với DOI: {doi}")
                return doi, crossref_info
            else:
                logging.warning(f"Không tìm thấy DOI cho reference: {reference}")
                return None, {}
        else:
            logging.error(f"Lỗi khi truy cập Crossref API: {response.status_code}")
            return None, {}
    except Exception as e:
        logging.error(f"Lỗi khi tìm kiếm DOI cho reference '{reference}': {e}")
        return None, {}
    finally:
        time.sleep(1)

def download_paper_pdf(doi, download_dir="downloaded_papers", email=EMAIL):
    """
    Tải xuống tệp PDF của bài báo sử dụng DOI và Unpaywall API.

    Args:
        doi (str): DOI của bài báo.
        download_dir (str): Thư mục để lưu trữ tệp PDF.
        email (str): Địa chỉ email của bạn để sử dụng với Unpaywall API.

    Returns:
        str: Đường dẫn tới tệp PDF đã tải xuống, hoặc None nếu không thành công.
    """
    if not os.path.exists(download_dir):
        os.makedirs(download_dir)

    api_url = f"https://api.unpaywall.org/v2/{doi}?email={email}"

    try:
        response = requests.get(api_url, timeout=10)
        response.raise_for_status()
        data = response.json()

        best_oa_location = data.get('best_oa_location', None)
        if best_oa_location is None or best_oa_location.get('url_for_pdf') is None:
            logging.warning(f"Không có URL PDF mở cho DOI: {doi}")
            return None

        pdf_url = best_oa_location.get('url_for_pdf')
        title = data.get('title', 'unknown_title').replace('/', '_').replace('\\', '_')
        pdf_path = os.path.join(download_dir, f"{title}.pdf")

        try:
            headers = {"User-Agent": "Mozilla/5.0"}
            pdf_response = requests.get(pdf_url, headers=headers, timeout=10)
            pdf_response.raise_for_status()

            with open(pdf_path, 'wb') as f:
                f.write(pdf_response.content)
            logging.info(f"Đã tải xuống PDF: {pdf_path}")
            return pdf_path
        except Exception as e:
            logging.error(f"Lỗi khi tải xuống PDF từ {pdf_url}: {e}")
            return None

    except requests.exceptions.HTTPError as http_err:
        logging.error(f"Lỗi HTTP khi truy cập Unpaywall API: {http_err}")
    except Exception as e:
        logging.error(f"Lỗi khi sử dụng Unpaywall API với DOI '{doi}': {e}")
    finally:
        time.sleep(1)

    return None

def crawl_references(references_list, download_dir="downloaded_papers"):
    """
    Crawl các tài liệu tham khảo để lấy DOI và thông tin chi tiết.

    Args:
        references_list (list): Danh sách các tham chiếu.
        download_dir (str): Thư mục để lưu trữ tệp PDF.

    Returns:
        dict: Thông tin crawl cho mỗi tham chiếu.
    """
    email = EMAIL

    total_entries = len(references_list)
    logging.info(f"Tổng số tài liệu tham khảo cần xử lý: {total_entries}")

    reference_crawl_info = {}

    for idx, entry in enumerate(references_list, start=1):
        reference = entry.get('reference', '')
        logging.info(f"Đang xử lý tài liệu tham khảo {idx}/{total_entries}: {reference}")
        if not reference or reference == "Reference not found.":
            reference_crawl_info[reference] = {
                'doi': None,
                'crossref_info': {},
                'pdf_path': None
            }
            continue

        doi, crossref_info = search_paper_doi_with_info(reference)
        if not doi:
            reference_crawl_info[reference] = {
                'doi': None,
                'crossref_info': {},
                'pdf_path': None
            }
            continue

        pdf_path = download_paper_pdf(doi, download_dir, email)
        if not pdf_path:
            reference_crawl_info[reference] = {
                'doi': doi,
                'crossref_info': crossref_info,
                'pdf_path': None
            }
            continue

        reference_crawl_info[reference] = {
            'doi': doi,
            'crossref_info': crossref_info,
            'pdf_path': pdf_path
        }

        time.sleep(5)  # Chờ giữa các yêu cầu

    return reference_crawl_info

def generate_json_output(citation_entries, output_file):
    """
    Tạo tệp JSON từ danh sách các mục trích dẫn.

    Args:
        citation_entries (list): Danh sách các mục trích dẫn.
        output_file (str): Đường dẫn tới tệp JSON sẽ được tạo.
    """
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(citation_entries, f, ensure_ascii=False, indent=4)
    logging.info(f"Đã tạo tệp JSON tại: {output_file}")

def check_docx_existence(docx_file):
    """
    Kiểm tra xem tệp DOCX đã tồn tại hay chưa.

    Args:
        docx_file (str): Đường dẫn tới tệp DOCX.

    Returns:
        bool: True nếu tệp tồn tại, ngược lại là False.
    """
    return os.path.exists(docx_file)
