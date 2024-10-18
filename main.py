import spacy
import logging
from docx import Document
from helper import (
    convert_pdf_to_docx,
    detect_citation_format,
    clean_text,
    crawl_references,
    generate_json_output
)
from apa_module import extract_apa_citations_with_context
import ieee_module  # Import module IEEE
import json

def extract_sentences_from_docx(docx_file, output_file):
    """
    Sử dụng spaCy để lấy tất cả các câu từ file DOCX và in chúng ra file văn bản,
    đồng thời trả về một mảng các câu.

    Args:
        docx_file (str): Đường dẫn đến tệp DOCX.
        output_file (str): Đường dẫn đến tệp văn bản để lưu các câu.

    Returns:
        List[str]: Danh sách các câu trích xuất từ tài liệu.
    """
    logging.info(f"Đang đọc tệp DOCX: {docx_file}")

    # Load mô hình spaCy
    nlp = spacy.load("en_core_web_sm")

    # Đọc nội dung từ file DOCX
    doc = Document(docx_file)

    text = "\n".join([para.text for para in doc.paragraphs])
    cleaned_text = clean_text(text)
    # Xử lý văn bản bằng spaCy
    spacy_doc = nlp(cleaned_text)

    # Tạo danh sách các câu
    sentences = []

    # Ghi các câu vào file và lưu trong mảng
    with open(output_file, "w", encoding="utf-8") as f:
        for sent in spacy_doc.sents:
            cleaned_sentence = sent.text.strip()
            if cleaned_sentence.lower() == 'references':
                break
            f.write(cleaned_sentence + "\n")
            sentences.append(cleaned_sentence)  # Thêm câu vào danh sách

    logging.info(f"Đã lưu các câu vào tệp: {output_file}")
    return sentences  # Trả về danh sách các câu

def main():
    pdf_file = "extraction_ieee.pdf"  # Thay thế bằng đường dẫn tệp PDF của bạn
    docx_file = "paper.docx"
    output_sentences_file = "paper_spacy_sentences.txt"

    try:
        # Chuyển đổi PDF sang DOCX nếu cần
        # docx_file_path = convert_pdf_to_docx(pdf_file, docx_file)

        # Trích xuất câu từ DOCX
        sentences = extract_sentences_from_docx(docx_file, output_sentences_file)
        doc = Document(docx_file)

        text = "\n".join([para.text for para in doc.paragraphs])
        citation_type = detect_citation_format(text)
        if citation_type == "APA":
            citations = []
            for sentence in sentences:
                sentence_citations = extract_apa_citations_with_context(sentence)
                citations.extend(sentence_citations)
            generate_json_output(citations, "apa_output.json")

        elif citation_type == "IEEE":
            # Trích xuất references từ DOCX
            references_map = ieee_module.extract_references(docx_file)
            # Trích xuất các trích dẫn IEEE
            ieee_citations = ieee_module.extract_ieee_citations_with_context(sentences, references_map)

            # Crawl và lấy đường dẫn PDF cho các references
            verified_citations = crawl_references(ieee_citations)

            # Tạo file JSON với thông tin trích dẫn và đường dẫn PDF
            generate_json_output(verified_citations, "ieee_output.json")

        elif citation_type == "Chicago":
            print("Định dạng trích dẫn Chicago chưa được hỗ trợ.")
        else:
            print("Không nhận diện được định dạng trích dẫn.")
    except Exception as e:
        logging.error(f"Có lỗi trong quá trình thực thi: {e}")

if __name__ == "__main__":
    main()
